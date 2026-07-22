#!/usr/bin/env python3
from __future__ import annotations
import argparse, csv, json, re, sys
from dataclasses import asdict, dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

DEFAULT_WEEK_1_MONDAY='2026-09-21'
EXPECTED_WEEKS=12
MAX_DETAIL_CHARS=70
MAX_ASSESSMENT_CHARS=65
MAX_MLO_DESC_CHARS=220
PROJECT_ROOT=Path(__file__).resolve().parent.parent
DEFAULT_EASTER_CONFIG=PROJECT_ROOT/'config'/'easter_sunday_dates_2027_2036.csv'

@dataclass
class WeekExtract:
    week:int; date_label:str; title:str; detail:str; assessment:str; render_pill:bool

@dataclass
class MLOExtract:
    code:str; title:str; description:str

def ordinal(n:int)->str:
    if 10 <= n%100 <= 20: s='th'
    else: s={1:'st',2:'nd',3:'rd'}.get(n%10,'th')
    return f'{n}{s}'

def month_label(d:date)->str:
    return {1:'Jan',2:'Feb',3:'Mar',4:'Apr',5:'May',6:'Jun',7:'Jul',8:'Aug',9:'Sept',10:'Oct',11:'Nov',12:'Dec'}[d.month]

def format_date_range(start:date, end:date)->str:
    if start.month==end.month:
        return f'Week ({ordinal(start.day)} - {ordinal(end.day)} {month_label(end)})'
    return f'Week ({ordinal(start.day)} {month_label(start)} - {ordinal(end.day)} {month_label(end)})'

def format_break_range(start:date, end:date)->str:
    if start.month==end.month:
        return f'({ordinal(start.day)} - {ordinal(end.day)} {month_label(end)})'
    return f'({ordinal(start.day)} {month_label(start)} - {ordinal(end.day)} {month_label(end)})'

def load_easter_sundays(csv_path:Path)->Dict[int,date]:
    result:Dict[int,date]={}
    with csv_path.open(newline='', encoding='utf-8') as handle:
        rows=list(csv.reader(handle))
    for row in rows[1:]:
        if not row or not row[0].strip(): continue
        d=datetime.strptime(row[0].strip(), '%d/%m/%Y').date()
        result[d.year]=d
    return result

# Easter Sunday always falls within this window (22 Mar - 25 Apr) for any given year,
# regardless of the exact date. Used to skip requiring table coverage for years a term
# never gets close to, so an ordinary autumn term doesn't need next year's row filled in.
EASTER_WINDOW_EARLIEST=(3,22)
EASTER_WINDOW_LATEST=(4,25)

def compute_week_dates(weeks:List['WeekExtract'], week_1_monday:date, easter_dates:Dict[int,date])->Tuple[List[dict], Optional[Tuple[date,date]]]:
    n=len(weeks)
    if n==0: return [], None

    naive_start=week_1_monday
    naive_end=week_1_monday+timedelta(days=(n-1)*7+4)

    candidate_years=set()
    for year in {naive_start.year, naive_end.year}:
        window_start=date(year, *EASTER_WINDOW_EARLIEST)
        window_end=date(year, *EASTER_WINDOW_LATEST)
        if naive_start<=window_end and naive_end>=window_start:
            candidate_years.add(year)

    missing=sorted(y for y in candidate_years if y not in easter_dates)
    if missing:
        years_text=', '.join(str(y) for y in missing)
        raise ValueError(
            f'No Easter Sunday date found for year(s) {years_text} in the Easter '
            f'lookup table ({DEFAULT_EASTER_CONFIG.name}). Add the missing year(s) '
            'before generating a term that spans them.'
        )

    break_window:Optional[Tuple[date,date]]=None
    for year in sorted(candidate_years):
        easter_sunday=easter_dates[year]
        if naive_start<=easter_sunday<=naive_end:
            break_window=(easter_sunday-timedelta(days=6), easter_sunday+timedelta(days=5))
            break

    def make_break_entry(window:Tuple[date,date])->dict:
        break_range=format_break_range(window[0], window[1])
        return {'week':None,'date_label':break_range,'title':'','detail':'','assessment':f'Easter Break\n{break_range}','render_pill':True,'kind':'break'}

    result:List[dict]=[]
    offset_days=0
    break_inserted=False
    for w in weeks:
        start=week_1_monday+timedelta(days=(w.week-1)*7+offset_days)
        if break_window and not break_inserted and start>=break_window[0]:
            result.append(make_break_entry(break_window))
            offset_days+=14
            break_inserted=True
            start=week_1_monday+timedelta(days=(w.week-1)*7+offset_days)

        end=start+timedelta(days=4)
        entry=asdict(w)
        entry['date_label']=format_date_range(start, end)
        entry['kind']='week'
        result.append(entry)

    if break_window and not break_inserted:
        result.append(make_break_entry(break_window))

    return result, break_window

def clean_text(text:str)->str:
    return re.sub(r'\s+', ' ', text.replace('\u00a0',' ')).strip()

def is_bullet_paragraph(paragraph)->bool:
    style=(paragraph.style.name or '').lower() if paragraph.style else ''
    text=clean_text(paragraph.text)
    return 'bullet' in style or bool(re.match(r'^[•\-*]\s+', text))

def paragraph_is_boldish(paragraph)->bool:
    runs=[r for r in paragraph.runs if r.text.strip()]
    return bool(runs) and any(r.bold for r in runs)

def extract_title_and_detail(cell)->tuple[str,str]:
    non_empty=[p for p in cell.paragraphs if clean_text(p.text)]
    if not non_empty: return '', ''
    title_para=next((p for p in non_empty if paragraph_is_boldish(p)), non_empty[0])
    title=clean_text(title_para.text)
    detail=''; title_seen=False
    for p in non_empty:
        if p is title_para:
            title_seen=True; continue
        if not title_seen: continue
        if is_bullet_paragraph(p):
            detail=re.sub(r'^[•\-*]\s+', '', clean_text(p.text)).strip(); break
    if not detail:
        for p in non_empty:
            if p is not title_para:
                detail=re.sub(r'^[•\-*]\s+', '', clean_text(p.text)).strip(); break
    return title, detail

def extract_assessment(cell)->str:
    return ' '.join(clean_text(p.text) for p in cell.paragraphs if clean_text(p.text)).strip()

def normalise_header(text:str)->str:
    return re.sub(r'\s+', ' ', text.strip().lower())

def find_week_table(doc):
    for table in doc.tables:
        if len(table.rows)<2: continue
        headers=[normalise_header(c.text) for c in table.rows[0].cells]
        if any(h=='week' for h in headers) and any('title' in h and 'topics' in h for h in headers) and any('assessment' in h for h in headers):
            return table
    return None

def header_index_map(table)->Dict[str,int]:
    headers=[normalise_header(c.text) for c in table.rows[0].cells]
    m={}
    for i,h in enumerate(headers):
        if h=='week': m['week']=i
        elif 'title' in h and 'topics' in h: m['title_topics']=i
        elif 'assessment' in h: m['assessments']=i
    return m

def extract_module_title(doc)->str:
    for table in doc.tables:
        if len(table.columns)<2: continue
        for row in table.rows:
            if clean_text(row.cells[0].text).lower()=='module title':
                return clean_text(row.cells[1].text)
    return 'Learner Journey Map'

def extract_mlo_cell_text(doc)->str:
    for table in doc.tables:
        if len(table.columns)<2:
            continue
        for row in table.rows:
            if clean_text(row.cells[0].text).lower()=='module learning outcomes':
                return '\n'.join(clean_text(p.text) for p in row.cells[1].paragraphs if clean_text(p.text)).strip()
    return ''

def extract_module_learning_outcomes(doc)->List[MLOExtract]:
    raw=extract_mlo_cell_text(doc)
    if not raw:
        return []

    # Typical source format: "Communication (MLO5): ..."
    pattern=re.compile(r'([^\n:]+?)\s*\((MLO\d+)\)\s*:\s*', re.IGNORECASE)
    matches=list(pattern.finditer(raw))
    outcomes:List[MLOExtract]=[]
    if matches:
        for i,m in enumerate(matches):
            start=m.end()
            end=matches[i+1].start() if i+1<len(matches) else len(raw)
            title=clean_text(m.group(1))
            code=m.group(2).upper()
            desc=clean_text(raw[start:end])
            outcomes.append(MLOExtract(code, title, desc))
        return outcomes

    # Fallback: split lines and use the first colon as divider.
    for line in [clean_text(x) for x in raw.split('\n') if clean_text(x)]:
        code_match=re.search(r'(MLO\d+)', line, re.IGNORECASE)
        code=code_match.group(1).upper() if code_match else 'MLO?'
        if ':' in line:
            left,right=line.split(':',1)
            left=re.sub(r'\(MLO\d+\)', '', left, flags=re.IGNORECASE).strip()
            title=left or code
            desc=clean_text(right)
        else:
            title=code
            desc=line
        outcomes.append(MLOExtract(code,title,desc))
    return outcomes

def extract_weeks(docx_path:Path):
    # python-docx only runs its own "is this actually a docx" check when given a str path —
    # passed a Path object, it skips straight to opening it as a zip and lets a raw, unhelpful
    # zipfile.BadZipFile escape instead of the library's own PackageNotFoundError.
    doc=Document(str(docx_path)); module_title=extract_module_title(doc); table=find_week_table(doc)
    if table is None: raise ValueError('Could not find week table')
    idx=header_index_map(table); weeks=[]
    for row in table.rows[1:]:
        wt=clean_text(row.cells[idx['week']].text); m=re.search(r'\d+', wt)
        if not m: continue
        n=int(m.group(0)); title,detail=extract_title_and_detail(row.cells[idx['title_topics']]); assessment=extract_assessment(row.cells[idx['assessments']])
        weeks.append(WeekExtract(n, '', title, detail, assessment, bool(assessment.strip())))
    mlos=extract_module_learning_outcomes(doc)
    return module_title, sorted(weeks, key=lambda w:w.week), mlos

def validate_weeks(weeks, expected):
    issues=[]
    if len(weeks)!=expected: issues.append(f'Expected {expected} week rows but found {len(weeks)}.')
    for w in weeks:
        if not w.title: issues.append(f'Week {w.week}: missing title.')
        if not w.detail: issues.append(f'Week {w.week}: missing detail.')
        if len(w.detail)>MAX_DETAIL_CHARS: issues.append(f'Week {w.week}: detail is {len(w.detail)} characters; consider shortening for the poster.')
        if w.render_pill and len(w.assessment)>MAX_ASSESSMENT_CHARS: issues.append(f'Week {w.week}: assessment pill text is {len(w.assessment)} characters; consider shortening.')
    return issues

def validate_mlos(mlos:List[MLOExtract]):
    issues=[]
    if not mlos:
        issues.append('No Module Learning Outcomes found in the header table.')
        return issues
    if len(mlos)<4:
        issues.append(f'Only {len(mlos)} MLO entries found; expected at least 4 for square render.')
    for mlo in mlos:
        if not mlo.code or not re.match(r'^MLO\d+$', mlo.code):
            issues.append(f'Unclear MLO code for entry "{mlo.title}".')
        if not mlo.description:
            issues.append(f'{mlo.code}: missing description.')
        if len(mlo.description)>MAX_MLO_DESC_CHARS:
            issues.append(f'{mlo.code}: description is {len(mlo.description)} characters; may wrap heavily in MLO square render.')
    return issues

def status(issues): return 'APPROVED' if not issues else 'NEEDS REVIEW'

def fail(message:str)->None:
    print(f'[FAIL] {message}', file=sys.stderr)
    sys.exit(1)

def main():
    ap=argparse.ArgumentParser(); ap.add_argument('--input', required=True); ap.add_argument('--review', required=True); ap.add_argument('--json', required=True); ap.add_argument('--week1', default=DEFAULT_WEEK_1_MONDAY); ap.add_argument('--expected-weeks', type=int, default=EXPECTED_WEEKS); ap.add_argument('--easter-config', default=str(DEFAULT_EASTER_CONFIG))
    a=ap.parse_args()
    input_path=Path(a.input)

    try:
        week1=datetime.strptime(a.week1,'%Y-%m-%d').date()
    except ValueError:
        fail(f'Term start date "{a.week1}" is not valid — expected format YYYY-MM-DD.')

    try:
        module,weeks,mlos=extract_weeks(input_path)
    except PackageNotFoundError:
        fail(f'"{input_path.name}" doesn\'t look like a valid Word (.docx) file. Please check the file and try again.')
    except ValueError as exc:
        if 'week table' in str(exc).lower():
            fail(f'Could not find a weekly plan table in "{input_path.name}". Make sure the document has a table with Week, Title & Topics, and Assessments columns.')
        else:
            fail(str(exc))

    issues=validate_weeks(weeks,a.expected_weeks)+validate_mlos(mlos)
    easter_dates=load_easter_sundays(Path(a.easter_config))
    try:
        final_weeks,break_window=compute_week_dates(weeks, week1, easter_dates)
    except ValueError as exc:
        fail(str(exc))
    easter_break_payload={'start':break_window[0].isoformat(),'end':break_window[1].isoformat()} if break_window else None
    payload={'module_title':module,'week_1_monday':week1.isoformat(),'expected_weeks':a.expected_weeks,'weeks_found':len(weeks),'mlos_found':len(mlos),'extraction_status':status(issues),'issues':issues,'weeks':final_weeks,'easter_break':easter_break_payload,'mlos':[asdict(m) for m in mlos]}
    Path(a.json).write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding='utf-8')
    easter_break_line=f'{break_window[0].strftime("%d/%m/%Y")} - {break_window[1].strftime("%d/%m/%Y")} (inserted)' if break_window else 'none (term does not cover Easter)'
    lines=[f'{module} — Student Journey Map Extraction Review','='*72,'',f'Extraction status : {status(issues)}',f'Weeks found       : {len(weeks)}',f'Expected weeks    : {a.expected_weeks}',f'MLOs found        : {len(mlos)}',f'Week 1 Monday     : {week1.strftime("%d/%m/%Y")}','Date rule         : Monday to Friday',f'Easter break      : {easter_break_line}','','Validation issues','-----------------']
    lines += [f'- {i}' for i in issues] if issues else ['None']
    lines += ['','Extracted week data','-------------------','']
    for w in final_weeks:
        if w['kind']=='break':
            lines += ['Easter Break',f'  Date       : {easter_break_line}','']
        else:
            lines += [f'Week {w["week"]}',f'  Date       : {w["date_label"]}',f'  Title      : {w["title"]}',f'  Detail     : {w["detail"]}',f'  Assessment : {w["assessment"]}',f'  Render pill: {w["render_pill"]}','']
    lines += ['','Extracted MLO data','------------------','']
    for m in mlos:
        lines += [f'{m.code}',f'  Title       : {m.title}',f'  Description : {m.description}','']
    Path(a.review).write_text('\n'.join(lines), encoding='utf-8')
    print(f'[OK] Extracted {len(weeks)} weeks; status {status(issues)}')
if __name__=='__main__': main()
